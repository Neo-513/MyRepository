from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
import docx
import os
import xlrd

path = "D:/dzr/客户、供应商走访名单（按地区）"  # word文件夹路径
excel = "D:/dzr/天纺标2019-2021其他业务收入明细(1).xlsx"  # excel文件路径
home = "D:/dzr/save"  # 保存目录


def str2float(string):  # 含逗号的字符串数据转为小数
	return float(string.replace(",", ""))


def get_files():  # 获取文件名
	files = []  # 文件路径列表
	for folder in os.listdir(path.rstrip("/")):  # 遍历父文件夹
		if folder not in ["福建", "广东", "山东"]:
			files.extend([f"{path}/{folder}/{file}" for file in os.listdir(f"{path}/{folder}") if "docx" in file])
	return files


def get_datas():
	tables = xlrd.open_workbook(excel)  # 读excel所有页数据
	pages = ["2021", "2020", "2019"]  # 年份和word中的列数对应关系
	dic = {}
	for page in pages:
		table = tables.sheet_by_name(page)  # 读excel单页数据
		headers = table.row_values(0)  # 表头
		idxs = [headers.index(header) for header in headers if "含税" in header]  # 含税列索引

		for row in range(table.nrows):
			r = table.row_values(row)  # 单行数据（返回数组）
			company = r[0]  # 公司
			dic[company] = []
			for i in range(len(r)):
				if i in idxs and r[i] != "" and r[i] != 0:
					lst = [headers[i], "", "", ""]
					lst[pages.index(page) + 1] = r[i]
					dic[company].append(lst)
	return dic


def process(word):
	file = docx.Document(word)  # 读word文件
	company = word.split("/")[-1][:-5].strip()  # 公司名称

	table = file.tables[2]  # 获取word第三张表格
	header = ["", "2021年", "2020年", "2019年"]  # 表头行
	count = ["合计", table.cell(0, 3).text, table.cell(0, 4).text, table.cell(0, 5).text]  # 合计行
	data = datas[company] if company in datas else []  # 中间数据
	check = ["检测业务", str2float(count[1]), str2float(count[2]), str2float(count[3])]  # 检测业务行
	for d in data:
		for i in range(1, 4):
			if d[i]:
				check[i] -= d[i]

	results = [header, check]
	if data:
		results.extend(data)
	results.append(count)

	table.add_column(Cm(1))  # word添加列
	for i in range(len(results)):
		if i > 1:
			table.add_row()  # 添加行
		for j in range(4):
			idx = i * 4 + j  # 生成索引
			table.cell(0, idx).width = Inches(2.5)  # 设置列宽（总列宽为10）
			value = results[i][j]  # 取所需值
			if isinstance(value, float):
				value = f"{results[i][j]:,.2f}"  # 格式转换

			table.cell(0, idx).text = ""  # 重置值
			run = table.cell(0, idx).paragraphs[0].add_run(value)  # word添加值
			run.font.name = "楷体"  # 修改ascii码中的字体
			run.font.size = Pt(12)  # 设置字号
			run.font.bold = (i == 0)  # 设置加粗
			r = run._element  # 访问内置元素
			r.rPr.rFonts.set(qn('w:eastAsia'), "楷体")  # 修改中文字体

	folder = f"{home}/{word.split('/')[-2]}"
	if not os.path.exists(folder):
		os.mkdir(folder)  # 新建文件夹

	file_name = f"{folder}/{word.split('/')[-1]}"  # 文件名
	file.save(file_name)  # 保存文件
	print(f"[SAVE]    {file_name}")


if __name__ == "__main__":
	datas = get_datas()
	for f in get_files():
		process(f)
